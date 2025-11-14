import asyncio
import tempfile
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
import pdf2image
from app.schemas.document import NormalizedDocument
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Handles parsing of different document formats (PDF, DOCX, images) into a normalized format
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.jpg', '.jpeg', '.png', '.txt'}
    
    async def parse_document(self, file_path: str, document_id: str) -> NormalizedDocument:
        """
        Parse a document and return normalized content
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract metadata
        metadata = await self._extract_metadata(file_path, document_id)
        
        # Parse content based on file type
        if file_extension == '.pdf':
            sections = await self._parse_pdf(file_path)
        elif file_extension == '.docx':
            sections = await self._parse_docx(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            sections = await self._parse_image(file_path)
        elif file_extension == '.txt':
            sections = await self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        return NormalizedDocument(
            document_id=document_id,
            metadata=metadata,
            sections=sections
        )
    
    async def _extract_metadata(self, file_path: str, document_id: str) -> Dict[str, Any]:
        """
        Extract basic metadata from the file
        """
        file_path_obj = Path(file_path)
        stat = file_path_obj.stat()
        
        metadata = {
            "document_id": document_id,
            "title": file_path_obj.stem,
            "file_path": str(file_path_obj),
            "file_size": stat.st_size,
            "file_type": file_path_obj.suffix.lower(),
            "creation_date": stat.st_ctime,
            "modification_date": stat.st_mtime,
            "page_count": 0,
            "word_count": 0
        }
        
        return metadata
    
    async def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse PDF document into sections
        """
        try:
            reader = PdfReader(file_path)
            sections = []
            word_count = 0
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text.strip():  # Only add sections with content
                    # Simple heuristic to identify headings (all caps, short lines)
                    lines = text.split('\n')
                    current_section = {
                        "type": "paragraph",
                        "level": 1,
                        "content": text,
                        "position": {"page": page_num, "order": len(sections) + 1}
                    }
                    
                    # Count words in the section
                    word_count += len(text.split())
                    
                    sections.append(current_section)
            
            # Update metadata with page count and word count
            # This would need to be done in the calling function
            
            return sections
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    async def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX document into sections
        """
        try:
            doc = DocxDocument(file_path)
            sections = []
            word_count = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Determine if it's a heading based on style
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.replace('Heading ', '').replace(' heading', ''))
                        section_type = "heading"
                    else:
                        level = 1
                        section_type = "paragraph"
                    
                    current_section = {
                        "type": section_type,
                        "level": level,
                        "content": paragraph.text,
                        "position": {"page": 1, "order": len(sections) + 1}  # DOCX doesn't have real pages
                    }
                    
                    # Count words in the section
                    word_count += len(paragraph.text.split())
                    
                    sections.append(current_section)
            
            # Process tables
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if cell.text.strip():
                            current_section = {
                                "type": "table",
                                "level": 1,
                                "content": cell.text,
                                "position": {"page": 1, "order": len(sections) + 1}
                            }
                            word_count += len(cell.text.split())
                            sections.append(current_section)
            
            return sections
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    async def _parse_image(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse image document using OCR
        """
        try:
            # Use PIL to open the image
            image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                sections = [{
                    "type": "paragraph",
                    "level": 1,
                    "content": text,
                    "position": {"page": 1, "order": 1}
                }]
                return sections
            else:
                # If OCR fails, try converting PDF to images first (for scanned PDFs)
                # This is a simplified approach; in reality, we'd need to detect if it's a scanned PDF
                return []
        except Exception as e:
            logger.error(f"Error parsing image {file_path}: {str(e)}")
            raise
    
    async def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse plain text document
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            sections = []
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    sections.append({
                        "type": "paragraph",
                        "level": 1,
                        "content": paragraph.strip(),
                        "position": {"page": 1, "order": i + 1}
                    })
            
            return sections
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise


# Enhanced version using PyMuPDF for better PDF handling
class EnhancedDocumentParser:
    """
    Enhanced document parser using PyMuPDF for better PDF handling and OCR capabilities
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.jpg', '.jpeg', '.png', '.txt', '.pptx', '.ppt'}
    
    async def parse_document(self, file_path: str, document_id: str) -> NormalizedDocument:
        """
        Parse a document and return normalized content using enhanced methods
        """
        import fitz  # PyMuPDF
        from docx import Document as DocxDocument
        import pytesseract
        from PIL import Image
        import pdf2image
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract metadata
        metadata = await self._extract_metadata(file_path, document_id)
        
        # Parse content based on file type
        if file_extension == '.pdf':
            sections = await self._parse_pdf_enhanced(file_path)
        elif file_extension == '.docx':
            sections = await self._parse_docx(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            sections = await self._parse_image_enhanced(file_path)
        elif file_extension == '.txt':
            sections = await self._parse_txt(file_path)
        else:
            # For other formats, we might use pandoc or other converters
            raise ValueError(f"Format not yet implemented: {file_extension}")
        
        # Update metadata with page count and word count
        metadata["page_count"] = len([s for s in sections if s["position"]["page"]])
        metadata["word_count"] = sum(len(s["content"].split()) for s in sections if s["type"] == "paragraph")
        
        return NormalizedDocument(
            document_id=document_id,
            metadata=metadata,
            sections=sections
        )
    
    async def _extract_metadata(self, file_path: str, document_id: str) -> Dict[str, Any]:
        """
        Extract basic metadata from the file
        """
        file_path_obj = Path(file_path)
        stat = file_path_obj.stat()
        
        metadata = {
            "document_id": document_id,
            "title": file_path_obj.stem,
            "file_path": str(file_path_obj),
            "file_size": stat.st_size,
            "file_type": file_path_obj.suffix.lower(),
            "creation_date": stat.st_ctime,
            "modification_date": stat.st_mtime,
            "page_count": 0,
            "word_count": 0
        }
        
        return metadata
    
    async def _parse_pdf_enhanced(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Enhanced PDF parsing using PyMuPDF
        """
        import fitz  # PyMuPDF
        
        try:
            doc = fitz.open(file_path)
            sections = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text with formatting info
                text = page.get_text("text")
                
                if text.strip():
                    # Extract blocks to identify headings
                    blocks = page.get_text("dict")["blocks"]
                    
                    for block in blocks:
                        if "lines" in block:  # Text block
                            # Combine lines in the block
                            block_text = ""
                            for line in block["lines"]:
                                line_text = ""
                                for span in line["spans"]:
                                    line_text += span["text"]
                                block_text += line_text + "\n"
                            
                            if block_text.strip():
                                # Simple heuristic to identify headings based on font size
                                # This is a simplified approach - in reality, we'd analyze font properties more thoroughly
                                font_sizes = []
                                for line in block["lines"]:
                                    for span in line["spans"]:
                                        font_sizes.append(span["size"])
                                
                                avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
                                # Assume headings are larger than 14pt
                                is_heading = avg_font_size > 14
                                
                                current_section = {
                                    "type": "heading" if is_heading else "paragraph",
                                    "level": 1,  # Could be determined by font size hierarchy
                                    "content": block_text.strip(),
                                    "position": {"page": page_num + 1, "order": len(sections) + 1}
                                }
                                
                                sections.append(current_section)
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n < 5:  # Check if not CMYK
                        img_data = pix.tobytes("png")
                        current_section = {
                            "type": "image",
                            "level": 1,
                            "content": f"Image on page {page_num + 1}, index {img_index}",
                            "position": {"page": page_num + 1, "order": len(sections) + 1}
                        }
                        sections.append(current_section)
                    pix = None  # Free memory
            
            doc.close()
            return sections
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    async def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX document into sections
        """
        try:
            doc = DocxDocument(file_path)
            sections = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Determine if it's a heading based on style
                    if paragraph.style.name.startswith('Heading'):
                        try:
                            level = int(paragraph.style.name.replace('Heading ', '').replace(' heading', ''))
                            section_type = "heading"
                        except ValueError:
                            level = 1
                            section_type = "paragraph"
                    else:
                        level = 1
                        section_type = "paragraph"
                    
                    current_section = {
                        "type": section_type,
                        "level": level,
                        "content": paragraph.text,
                        "position": {"page": 1, "order": len(sections) + 1}  # DOCX doesn't have real pages
                    }
                    
                    sections.append(current_section)
            
            # Process tables
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if cell.text.strip():
                            current_section = {
                                "type": "table",
                                "level": 1,
                                "content": cell.text,
                                "position": {"page": 1, "order": len(sections) + 1}
                            }
                            sections.append(current_section)
            
            return sections
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    async def _parse_image_enhanced(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Enhanced image parsing with better OCR preprocessing
        """
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            
            # Open and preprocess image for better OCR
            image = Image.open(file_path)
            
            # Enhance image for better OCR results
            # Increase contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2)
            
            # Sharpen image
            image = image.filter(ImageFilter.SHARPEN)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                sections = [{
                    "type": "paragraph",
                    "level": 1,
                    "content": text,
                    "position": {"page": 1, "order": 1}
                }]
                return sections
            else:
                return []
        except Exception as e:
            logger.error(f"Error parsing image {file_path}: {str(e)}")
            raise
    
    async def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse plain text document
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            sections = []
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    sections.append({
                        "type": "paragraph",
                        "level": 1,
                        "content": paragraph.strip(),
                        "position": {"page": 1, "order": i + 1}
                    })
            
            return sections
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise


# Singleton instance
document_parser = EnhancedDocumentParser()